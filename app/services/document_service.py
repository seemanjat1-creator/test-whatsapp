import aiofiles
import os
from typing import List, Optional, Dict, Any
from fastapi import UploadFile, HTTPException
from app.config import settings
from app.models.document import (
    Document, DocumentCreate, DocumentType, DocumentStatus, 
    DocumentChunk, SearchResult, DocumentSearch
)
from app.services.openai_service import openai_service
from app.services.excel_processor import excel_processor
from app.database import get_database
from bson import ObjectId
import PyPDF2
import docx
import logging
from datetime import datetime
import hashlib
import re
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import asyncio
from concurrent.futures import ThreadPoolExecutor

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self):
        self.upload_dir = settings.upload_dir
        os.makedirs(self.upload_dir, exist_ok=True)
        self.chunk_size = 800  # Optimal chunk size for embeddings
        self.chunk_overlap = 100  # Overlap between chunks
        self.max_chunks_per_document = 100  # Prevent excessive chunking
        self.executor = ThreadPoolExecutor(max_workers=2)  # For CPU-intensive tasks
    
    async def upload_document(
        self,
        file: UploadFile,
        workspace_id: str,
        title: Optional[str] = None,
        description: Optional[str] = None,
        tags: List[str] = None
    ) -> Document:
        """Upload and process document with chunking and embeddings"""
        try:
            logger.info(f"Starting document upload: {file.filename}, size: {file.size}, workspace: {workspace_id}")
            
            # Validate inputs
            if not file or not file.filename:
                logger.error("No file provided")
                raise HTTPException(
                    status_code=400,
                    detail="No file provided"
                )
            
            if not workspace_id or not workspace_id.strip():
                logger.error("No workspace ID provided")
                raise HTTPException(
                    status_code=400,
                    detail="Workspace ID is required"
                )
            
            # Validate file type
            if not self._is_valid_file_type(file.filename):
                logger.error(f"Invalid file type: {file.filename}, content_type: {file.content_type}")
                raise HTTPException(
                    status_code=400,
                    detail="Invalid file type. Only PDF, DOCX, TXT, and Excel files (.xlsx, .xls) are allowed."
                )
            
            # Check file size
            if file.size and file.size > settings.max_file_size:
                logger.error(f"File too large: {file.size} bytes")
                raise HTTPException(
                    status_code=400,
                    detail=f"File size too large. Maximum size is {settings.max_file_size / (1024*1024):.1f}MB"
                )
            
            # Save file temporarily
            file_path = await self._save_file(file)
            logger.info(f"File saved temporarily at: {file_path}")
            
            # Extract text content
            content = await self._extract_text(file_path, file.filename)
            logger.info(f"Text extracted, length: {len(content)} characters")
            
            if not content or len(content.strip()) < 10:
                logger.error("Document content too short or empty")
                raise HTTPException(
                    status_code=400,
                    detail="Document content is too short or could not be extracted"
                )
            
            # Create document record
            document_data = DocumentCreate(
                title=(title.strip() if title else None) or file.filename,
                file_name=file.filename,
                document_type=self._get_document_type(file.filename),
                content=content,
                file_size=file.size or 0,
                status=DocumentStatus.PROCESSING,
                tags=[tag.strip() for tag in (tags or []) if tag.strip()],
                description=description.strip() if description else None,
                workspace_id=workspace_id
            )
            
            # Save to database
            db = get_database()
            doc_dict = document_data.dict()
            doc_dict["created_at"] = datetime.utcnow()
            doc_dict["updated_at"] = datetime.utcnow()
            doc_dict["workspace_id"] = ObjectId(workspace_id)
            doc_dict["access_count"] = 0
            
            result = await db.documents.insert_one(doc_dict)
            document_id = str(result.inserted_id)
            
            # Process document in background (chunking and embeddings)
            await self._process_document_chunks(document_id, content, workspace_id)
            
            # Update document status
            await db.documents.update_one(
                {"_id": ObjectId(document_id)},
                {"$set": {
                    "status": DocumentStatus.READY,
                    "processed_at": datetime.utcnow()
                }}
            )
            
            # Get updated document
            doc_dict["_id"] = document_id
            doc_dict["workspace_id"] = workspace_id
            doc_dict["status"] = DocumentStatus.READY
            doc_dict["processed_at"] = datetime.utcnow()
            
            return Document(**doc_dict)
            
        except Exception as e:
            logger.error(f"Document upload error: {e}")
            # Update status to error if document was created
            if 'document_id' in locals():
                db = get_database()
                await db.documents.update_one(
                    {"_id": ObjectId(document_id)},
                    {"$set": {"status": DocumentStatus.ERROR}}
                )
            raise HTTPException(status_code=500, detail="Failed to upload document")
    
    async def _process_document_chunks(self, document_id: str, content: str, workspace_id: str):
        """Process document into chunks with embeddings"""
        try:
            logger.info(f"Starting chunk processing for document {document_id}")
            db = get_database()
            
            # Get document to check type
            doc_data = await db.documents.find_one({"_id": ObjectId(document_id)})
            if not doc_data:
                raise ValueError(f"Document {document_id} not found")
            
            # Use Excel-specific chunking for Excel files
            if doc_data.get("document_type") in ["xlsx", "xls"]:
                chunk_data_list = excel_processor.create_excel_chunks(content, doc_data["file_name"])
                chunks = [chunk_data["content"] for chunk_data in chunk_data_list]
                chunk_metadata_list = [chunk_data["metadata"] for chunk_data in chunk_data_list]
            else:
                # Use standard chunking for other file types
                chunks = self._split_into_chunks(content)
                chunk_metadata_list = [{"chunk_type": "standard"} for _ in chunks]
            
            logger.info(f"Document split into {len(chunks)} chunks")
            
            # Limit chunks to prevent excessive processing
            if len(chunks) > self.max_chunks_per_document:
                chunks = chunks[:self.max_chunks_per_document]
                chunk_metadata_list = chunk_metadata_list[:self.max_chunks_per_document]
                logger.warning(f"Limited chunks to {self.max_chunks_per_document} for document {document_id}")
            
            # Generate embeddings for each chunk
            chunk_documents = []
            for i, (chunk_content, chunk_metadata) in enumerate(zip(chunks, chunk_metadata_list)):
                if chunk_content.strip():  # Skip empty chunks
                    logger.info(f"Generating embedding for chunk {i+1}/{len(chunks)}")
                    embedding = await openai_service.generate_embedding(chunk_content)
                    
                    if embedding:  # Only add if embedding was generated successfully
                        # Merge standard metadata with chunk-specific metadata
                        combined_metadata = {
                            "word_count": len(chunk_content.split()),
                            "char_count": len(chunk_content),
                            **chunk_metadata
                        }
                        
                        chunk_doc = {
                            "document_id": document_id,
                            "workspace_id": ObjectId(workspace_id),
                            "content": chunk_content,
                            "chunk_index": i,
                            "embedding": embedding,
                            "metadata": combined_metadata,
                            "created_at": datetime.utcnow()
                        }
                        chunk_documents.append(chunk_doc)
                    else:
                        logger.warning(f"Failed to generate embedding for chunk {i}")
            
            # Insert chunks in batch
            if chunk_documents:
                await db.document_chunks.insert_many(chunk_documents)
                logger.info(f"Inserted {len(chunk_documents)} chunks for document {document_id}")
            else:
                logger.warning(f"No chunks with embeddings created for document {document_id}")
            
            # Update document with chunk count
            await db.documents.update_one(
                {"_id": ObjectId(document_id)},
                {"$set": {"chunk_count": len(chunk_documents)}}
            )
            
        except Exception as e:
            logger.error(f"Document chunk processing error: {e}")
    
    def _split_into_chunks(self, content: str) -> List[str]:
        """Split content into overlapping chunks"""
        if not content or len(content.strip()) < 50:
            return []
            
        chunks = []
        start = 0
        content = content.strip()
        
        while start < len(content):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(content):
                # Look for sentence endings
                sentence_end = max(
                    content.rfind('.', start, end),
                    content.rfind('!', start, end),
                    content.rfind('?', start, end)
                )
                if sentence_end > start + self.chunk_size // 3:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    word_end = content.rfind(' ', start, end)
                    if word_end > start + self.chunk_size // 3:
                        end = word_end
            
            chunk = content[start:end].strip()
            if chunk and len(chunk) > 20:  # Only include meaningful chunks
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(content):
                break
                
            # Prevent infinite loops
            if len(chunks) > self.max_chunks_per_document:
                logger.warning(f"Reached maximum chunk limit: {self.max_chunks_per_document}")
                break
        
        return chunks
    
    async def get_workspace_documents(self, workspace_id: str) -> List[Document]:
        """Get all documents for a workspace with stats"""
        db = get_database()
        cursor = db.documents.find({"workspace_id": ObjectId(workspace_id)}).sort("created_at", -1)
        documents = []
        
        async for doc in cursor:
            doc["_id"] = str(doc["_id"])
            doc["workspace_id"] = workspace_id
            documents.append(Document(**doc))
        
        return documents
    
    async def get_document_by_id(self, document_id: str, workspace_id: str) -> Optional[Document]:
        """Get document by ID and update access stats"""
        db = get_database()
        
        # Update access stats
        await db.documents.update_one(
            {"_id": ObjectId(document_id), "workspace_id": ObjectId(workspace_id)},
            {
                "$inc": {"access_count": 1},
                "$set": {"last_accessed": datetime.utcnow()}
            }
        )
        
        doc_data = await db.documents.find_one({
            "_id": ObjectId(document_id),
            "workspace_id": ObjectId(workspace_id)
        })
        
        if not doc_data:
            return None
        
        doc_data["_id"] = str(doc_data["_id"])
        doc_data["workspace_id"] = workspace_id
        
        return Document(**doc_data)
    
    async def update_document(
        self, 
        document_id: str, 
        workspace_id: str, 
        update_data: Dict[str, Any]
    ) -> Optional[Document]:
        """Update document metadata"""
        db = get_database()
        
        update_data["updated_at"] = datetime.utcnow()
        
        result = await db.documents.update_one(
            {"_id": ObjectId(document_id), "workspace_id": ObjectId(workspace_id)},
            {"$set": update_data}
        )
        
        if result.modified_count == 0:
            return None
        
        return await self.get_document_by_id(document_id, workspace_id)
    
    async def delete_document(self, document_id: str, workspace_id: str) -> bool:
        """Delete a document and its chunks"""
        db = get_database()
        
        # Delete document chunks first
        await db.document_chunks.delete_many({
            "document_id": document_id,
            "workspace_id": ObjectId(workspace_id)
        })
        
        # Delete document
        result = await db.documents.delete_one({
            "_id": ObjectId(document_id),
            "workspace_id": ObjectId(workspace_id)
        })
        
        return result.deleted_count > 0
    
    async def search_documents(
        self,
        search_request: DocumentSearch
    ) -> List[SearchResult]:
        """Enhanced vector search with MongoDB Atlas Vector Search integration"""
        try:
            logger.info(f"Starting document search for query: '{search_request.query}' in workspace {search_request.workspace_id}")
            db = get_database()
            
            # Generate query embedding
            query_embedding = await openai_service.generate_embedding(search_request.query)
            
            if not query_embedding:
                logger.warning("Failed to generate query embedding")
                return []
            
            logger.info("Query embedding generated successfully")
            
            # MongoDB Atlas Vector Search pipeline (if available)
            # For now, using cosine similarity with existing chunks
            vector_search_pipeline = [
                {
                    "$match": {
                        "workspace_id": ObjectId(search_request.workspace_id),
                        "embedding": {"$exists": True, "$ne": None}
                    }
                },
                {
                    "$lookup": {
                        "from": "documents",
                        "localField": "document_id",
                        "foreignField": "_id",
                        "as": "document"
                    }
                },
                {
                    "$unwind": "$document"
                },
                {
                    "$match": {
                        "document.status": "ready"
                    }
                }
            ]
            
            # Add document type filter if specified
            if search_request.document_types:
                vector_search_pipeline.append({
                    "$match": {
                        "document.document_type": {"$in": search_request.document_types}
                    }
                })
            
            # Add tags filter if specified
            if search_request.tags:
                vector_search_pipeline.append({
                    "$match": {
                        "document.tags": {"$in": search_request.tags}
                    }
                })
            
            # Execute chunk search
            chunks = []
            chunk_count = 0
            async for chunk in db.document_chunks.aggregate(vector_search_pipeline):
                chunk_count += 1
                if chunk.get("embedding"):
                    try:
                        similarity = cosine_similarity(
                            [query_embedding],
                            [chunk["embedding"]]
                        )[0][0]
                    except Exception as e:
                        logger.warning(f"Failed to calculate similarity for chunk: {e}")
                        continue
                    
                    if similarity >= search_request.similarity_threshold:
                        chunk["similarity_score"] = float(similarity)
                        chunks.append(chunk)
            
            logger.info(f"Processed {chunk_count} chunks, found {len(chunks)} matching chunks")
            
            # Sort by similarity
            chunks.sort(key=lambda x: x["similarity_score"], reverse=True)
            
            # Group by document and calculate relevance
            document_results = {}
            for chunk in chunks[:search_request.limit * 5]:  # Get more chunks for better grouping
                doc_id = str(chunk["document"]["_id"])
                
                if doc_id not in document_results:
                    document_results[doc_id] = {
                        "document": chunk["document"],
                        "chunks": [],
                        "max_similarity": 0,
                        "avg_similarity": 0,
                        "total_similarity": 0,
                        "chunk_count": 0
                    }
                
                result = document_results[doc_id]
                result["chunks"].append({
                    "_id": str(chunk["_id"]),
                    "document_id": doc_id,
                    "workspace_id": search_request.workspace_id,
                    "content": chunk["content"],
                    "chunk_index": chunk["chunk_index"],
                    "embedding": chunk["embedding"],
                    "metadata": chunk.get("metadata", {}),
                    "created_at": chunk["created_at"]
                })
                
                result["max_similarity"] = max(result["max_similarity"], chunk["similarity_score"])
                result["total_similarity"] += chunk["similarity_score"]
                result["chunk_count"] += 1
                result["avg_similarity"] = result["total_similarity"] / result["chunk_count"]
            
            logger.info(f"Grouped results into {len(document_results)} documents")
            
            # Create search results
            search_results = []
            for doc_id, result in document_results.items():
                # Calculate relevance score (combination of max and average similarity)
                relevance_score = (result["max_similarity"] * 0.6) + (result["avg_similarity"] * 0.4)
                
                # Prepare document
                doc_data = result["document"]
                doc_data["_id"] = str(doc_data["_id"])
                doc_data["workspace_id"] = search_request.workspace_id
                
                search_result = SearchResult(
                    document=Document(**doc_data),
                    chunks=result["chunks"][:5],  # Top 5 chunks per document
                    similarity_score=result["max_similarity"],
                    relevance_score=relevance_score
                )
                search_results.append(search_result)
            
            # Sort by relevance and limit results
            search_results.sort(key=lambda x: x.relevance_score, reverse=True)
            final_results = search_results[:search_request.limit]
            
            logger.info(f"Returning {len(final_results)} search results")
            return final_results
            
        except Exception as e:
            logger.error(f"Document search error: {e}")
            return []
    
    async def get_document_stats(self, workspace_id: str) -> Dict[str, Any]:
        """Get document statistics for workspace"""
        db = get_database()
        
        pipeline = [
            {"$match": {"workspace_id": ObjectId(workspace_id)}},
            {
                "$group": {
                    "_id": None,
                    "total_documents": {"$sum": 1},
                    "total_size": {"$sum": "$file_size"},
                    "total_chunks": {"$sum": "$chunk_count"},
                    "avg_access_count": {"$avg": "$access_count"},
                    "document_types": {"$push": "$document_type"},
                    "statuses": {"$push": "$status"}
                }
            }
        ]
        
        result = await db.documents.aggregate(pipeline).to_list(1)
        
        if not result:
            return {
                "total_documents": 0,
                "total_size": 0,
                "total_chunks": 0,
                "avg_access_count": 0,
                "type_breakdown": {},
                "status_breakdown": {}
            }
        
        stats = result[0]
        
        # Calculate type breakdown
        type_breakdown = {}
        for doc_type in stats.get("document_types", []):
            type_breakdown[doc_type] = type_breakdown.get(doc_type, 0) + 1
        
        # Calculate status breakdown
        status_breakdown = {}
        for status in stats.get("statuses", []):
            status_breakdown[status] = status_breakdown.get(status, 0) + 1
        
        return {
            "total_documents": stats.get("total_documents", 0),
            "total_size": stats.get("total_size", 0),
            "total_chunks": stats.get("total_chunks", 0),
            "avg_access_count": round(stats.get("avg_access_count", 0), 2),
            "type_breakdown": type_breakdown,
            "status_breakdown": status_breakdown
        }
    
    async def _save_file(self, file: UploadFile) -> str:
        """Save uploaded file to disk"""
        # Create unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        file_hash = hashlib.md5(file.filename.encode()).hexdigest()[:8]
        unique_filename = f"{timestamp}_{file_hash}_{file.filename}"
        file_path = os.path.join(self.upload_dir, unique_filename)
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        return file_path
    
    async def _extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from different file types"""
        try:
            if filename.lower().endswith('.pdf'):
                return await self._extract_from_pdf(file_path)
            elif filename.lower().endswith('.docx'):
                return await self._extract_from_docx(file_path)
            elif filename.lower().endswith('.txt'):
                return await self._extract_from_txt(file_path)
            elif filename.lower().endswith(('.xlsx', '.xls')):
                return await self._extract_from_excel(file_path, filename)
            else:
                raise ValueError("Unsupported file type")
        finally:
            # Clean up file
            if os.path.exists(file_path):
                os.remove(file_path)
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # Clean up text
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = re.sub(r'\n+', '\n', text)  # Normalize line breaks
        return text.strip()
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        return content.strip()
    
    async def _extract_from_excel(self, file_path: str, filename: str) -> str:
        """Extract text from Excel files"""
        try:
            # Validate Excel file first
            is_valid, error_message = excel_processor.validate_excel_file(file_path, filename)
            if not is_valid:
                raise ValueError(error_message)
            
            # Process Excel file
            content = await excel_processor.process_excel_file(file_path, filename)
            
            if not content or len(content.strip()) < 10:
                raise ValueError("Excel file appears to be empty or contains no readable data")
            
            logger.info(f"Successfully extracted {len(content)} characters from Excel file {filename}")
            return content
            
        except Exception as e:
            logger.error(f"Excel extraction error for {filename}: {e}")
            raise ValueError(f"Failed to process Excel file: {str(e)}")
    
    def _is_valid_file_type(self, filename: str) -> bool:
        """Check if file type is valid"""
        valid_extensions = ['.pdf', '.docx', '.txt', '.xlsx', '.xls']
        return any(filename.lower().endswith(ext) for ext in valid_extensions)
    
    def _get_document_type(self, filename: str) -> DocumentType:
        """Get document type from filename"""
        if filename.lower().endswith('.pdf'):
            return DocumentType.PDF
        elif filename.lower().endswith('.docx'):
            return DocumentType.DOCX
        elif filename.lower().endswith('.txt'):
            return DocumentType.TXT
        elif filename.lower().endswith('.xlsx'):
            return DocumentType.XLSX
        elif filename.lower().endswith('.xls'):
            return DocumentType.XLS
        else:
            return DocumentType.TXT

document_service = DocumentService()